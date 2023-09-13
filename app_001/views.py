from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, HttpResponse
from django.shortcuts import render, redirect
from django.contrib import messages
from app_001.EmailBackEnd import EmailBackEnd
from django.views.decorators.csrf import csrf_exempt
from app_001 .models import *
from .forms import *
import docx
from docx import Document

def home(request):
    return render(request, 'index.html')

def loginPage(request):
    return render(request, 'login.html')

def course_reg_confirm(request):
        return render(request,'hod_template/course_registratin.html')

def course_reg_update(request):
    if request.method == "POST":
        aa=int(request.POST["radio"])
        if aa==1:
            ss=Status.objects.create(status=aa)
        elif aa==0:
            ss=Status.objects.all()
            ss.delete()    
        return render(request,'hod_template/course_registratin.html')
    
def doLogin(request):
    if request.method != "POST":
        return HttpResponse("<h2>## Error 404 ##</h2>")
    else:
        username=request.POST.get('username')
        password=request.POST.get('password')
        user = EmailBackEnd.authenticate(request,username, password)
        if user != None:
            login(request, user)
            user_type = user.user_type
            if user_type == '1':
                return redirect('admin_home')
                      
            elif user_type == '2':
                # return HttpResponse("Student Login")
                return redirect('student_home')
            else:
                messages.error(request, "Invalid Login!")
                return redirect('login')
        else:
            messages.error(request, "Invalid Login Credentials!")
            #return HttpResponseRedirect("/")
            return redirect('login')







def logout_user(request):
    logout(request)
    return HttpResponseRedirect('/')



############_____ADMIN______#############



def adminhome(request):
    all_student_count = Students.objects.all().count()
    #staff_count = Staffs.objects.all().count()
    all_student_leaves=LeaveReportStudent.objects.all().count()
    context={
        "all_student_count": all_student_count,  
        #"staff_count": staff_count,
        "all_student_leaves" : all_student_leaves
    }
    return render(request, "hod_template/home_content.html", context)


def addstudent(request):
    form = AddStudentForm()
    context = {
        "form": form
    }
    return render(request, 'hod_template/add_student_template.html', context)


def addstudentsave(request):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('add_student')
    else:
        form = AddStudentForm(request.POST)

        if form.is_valid():
            first_name = form.cleaned_data['first_name']
            last_name = form.cleaned_data['last_name']
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            rollno=form.cleaned_data['rollnumber']
            password = form.cleaned_data['password']
            address = form.cleaned_data['address']
            gender = form.cleaned_data['gender']
            user = CustomUser.objects.create_user(username=username, password=password, email=email,first_name=first_name, last_name=last_name, user_type=2)
            user.students.address = address
            user.students.rollnumber = rollno
            user.students.gender = gender
                
            user.save()
            messages.success(request, "Student Added Successfully!")
            return redirect('add_student')
        
        else:
            messages.error(request, "Failed to Add Student!")
            return redirect('add_student')


def managestudent(request):
    students = Students.objects.all()
    context = {
        "students": students
    }
    return render(request, 'hod_template/manage_student_template.html', context)


def editstudent(request, student_id):
    # Adding Student ID into Session Variable
    request.session['student_id'] = student_id

    student = Students.objects.get(admin=student_id)
    form = EditStudentForm()
    # Filling the form with Data from Database
    form.fields['email'].initial = student.admin.email
    form.fields['username'].initial = student.admin.username
    form.fields['rollnumber'].initial = student.rollnumber
    form.fields['first_name'].initial = student.admin.first_name
    form.fields['last_name'].initial = student.admin.last_name
    form.fields['address'].initial = student.address
    form.fields['gender'].initial = student.gender

    context = {
        "id": student_id,
        "username": student.admin.username,
        "form": form
    }
    return render(request, "hod_template/edit_student_template.html", context)


def edit_student_save(request):
    if request.method != "POST":
            return HttpResponse("Invalid Method!") 
     
    else:
            student_id = request.session.get('student_id')
            print('===========================================',student_id)
            if student_id == None:
                return redirect('/manage_student')
            form = EditStudentForm(request.POST)
            if form.is_valid():
                email = form.cleaned_data['email']
                username = form.cleaned_data['username']
                first_name = form.cleaned_data['first_name']
                last_name = form.cleaned_data['last_name']
                address = form.cleaned_data['address']
                gender = form.cleaned_data['gender']
                try:
                    user = CustomUser.objects.get(id=student_id)
                    user.first_name = first_name
                    user.last_name = last_name
                    user.email = email
                    user.username = username
                    user.save()
                    student_model = Students.objects.get(admin=student_id)
                    student_model.address = address
                    student_model.gender = gender
                    student_model.save()
                    del request.session['student_id']
                    messages.success(request, "Student Updated Successfully!")
                    return redirect('/edit_student/'+student_id)
                except:
                    messages.success(request, "Failed to Update Student.")
                    return redirect('/manage_student/'+student_id)
            else:
                return redirect('/manage_student/'+student_id)
   

def deletestudent(request, student_id):
    student = Students.objects.get(admin=student_id)
    try:
        student.delete()
        messages.success(request, "Student Deleted Successfully.")
        return redirect('manage_student')
    except:
        messages.error(request, "Failed to Delete Student.")
        return redirect('manage_student')


@csrf_exempt
def checkemailexist(request):
    email = request.POST.get("email")
    user_obj = CustomUser.objects.filter(email=email).exists()
    if user_obj:
        return HttpResponse(True)
    else:
        return HttpResponse(False)


@csrf_exempt
def checkusernameexist(request):
    username = request.POST.get("username")
    user_obj = CustomUser.objects.filter(username=username).exists()
    if user_obj:
        return HttpResponse(True)
    else:
        return HttpResponse(False)



##______STUDENT______##



def studenthome(request):
    student_obj = Students.objects.get(admin=request.user.id)
    return render(request, "student_template/student_home_template.html")


def studentapplyleave(request):
    student_obj = Students.objects.get(admin=request.user.id)
    leave_data = LeaveReportStudent.objects.filter(student_id=student_obj)
    context = {
        "leave_data": leave_data
    }
    return render(request, 'student_template/student_apply_leave.html', context)


def studentapplyleavesave(request):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('student_apply_leave')
    else:
        leavetype= request.POST['leave_type']
        fromleave_date = request.POST['fromleave_date']
        toleave_date = request.POST['toleave_date']
        leave_message = request.POST['leave_message']

        student_obj = Students.objects.get(admin=request.user.id)
        try:
            leave_report = LeaveReportStudent(student_id=student_obj, leave_types=leavetype, from_leave_date=fromleave_date,to_leave_date=toleave_date,leave_message=leave_message, leave_status=0)
            leave_report.save()
            messages.success(request, "Applied for Leave.")
            return redirect('student_apply_leave')
        except:
            messages.error(request, "Failed to Apply Leave")
            return redirect('student_apply_leave')
        
def cancelleave(request): 
    student_obj = Students.objects.get(admin=request.user.id)
    leave_deletes=LeaveReportStudent.objects.filter(student_id=student_obj)
    leave_deletes.delete()
    return redirect('student_apply_leave')



#########-LEAVE VIEWS-#################
def studentleaveview(request):
    leaves = LeaveReportStudent.objects.all()
    context = {
        "leaves": leaves
    }
    return render(request, 'hod_template/student_leave_view.html', context)

def studentleaveapprove(request, leave_id):
    leave = LeaveReportStudent.objects.get(id=leave_id)
    leave.leave_status = 1
    leave.save()
    return redirect('student_leave_view')


def studentleavereject(request, leave_id):
    leave = LeaveReportStudent.objects.get(id=leave_id)
    leave.leave_status = 2
    leave.save()
    return redirect('student_leave_view')


def student_register(request):
    aa=Status.objects.filter(status=1)
    bb=CouRegForm()
    return render(request,'student_template/course_reg_view.html',{'aa' : aa , 'bb':bb})

def crssavwe(request):
    student_obj = Students.objects.get(admin=request.user.id)  
    if request.method =="POST":
        stff=request.POST['stafflist']
        curs=request.POST['corselist']
        tms=request.POST['timess']
        aba=Courseregister.objects.create(Stafff=stff,Coursee=curs,timess=tms,studentid=student_obj)
        messages.success(request,'Course Registered Sucessfuly')
    return redirect('student_register')

def dowload_course_doc(request):
    document = Document()
    sampless=Students.objects.get(admin=request.user.id)
    document.add_heading('Course Registered', 0)

    p = document.add_paragraph('Time Table ')
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Course'
    hdr_cells[1].text = 'Staff'
    hdr_cells[2].text = 'Time'
    curs=Courseregister.objects.filter(studentid=sampless)
    print('@##################@@@@@@@@@@@@@@@@@@@@@@@@@@@##################@@@@@@@@@@@@@@@##########',curs)
    for i in curs:
        
        row = table.add_row().cells
        row[0].text = i.Coursee
        row[1].text = i.Stafff
        row[2].text = i.timess
        a=i.Stafff    
        print('@##################@@@@@@@@@@@@@@@@@@@@@@@@@@@##################@@@@@@@@@@@@@@@##########',a)
   
    #document.add_page_break()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{sampless.rollnumber}".docx'
    document.save(response)
    return response
 

